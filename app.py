import streamlit as st
import math
import re
from openai import OpenAI
import stripe
from docx import Document
from io import BytesIO

# 1. Page Configuration
st.set_page_config(
    page_title="Бизнес Навигатор",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Histats безплатен невидим брояч
st.markdown(
    '<a href="https://www.histats.com" target="_blank">'
    '<img src="https://sstatic1.histats.com/0.gif?5036919&101" alt="Histats" border="0" style="display:none;">'
    '</a>',
    unsafe_allow_html=True
)

# 🎨 ВИЗУАЛНА ТЕМА (тъмно червено-черно, зелени акценти, картови контейнери)
st.markdown(
    """
    <style>
    .stApp {
        background:
            radial-gradient(ellipse 90% 60% at 50% -5%, #6b0f0f 0%, transparent 60%),
            linear-gradient(180deg, #1c0505 0%, #100202 55%, #0a0101 100%);
    }
    .stApp::before {
        content: "";
        position: fixed;
        inset: 0;
        pointer-events: none;
        opacity: 0.05;
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='60' height='104' viewBox='0 0 60 104'%3E%3Cpolygon points='30,0 60,17 60,52 30,69 0,52 0,17' fill='none' stroke='%23ff6666' stroke-width='1'/%3E%3C/svg%3E");
        background-size: 60px 104px;
        z-index: 0;
    }

    /* --- Responsive ширина: пълен екран на десктоп, центрирано до разумен
       максимум, пълна ширина на мобилни устройства --- */
    [data-testid="stMainBlockContainer"] {
        position: relative;
        z-index: 1;
        max-width: 900px;
        margin: 0 auto;
        padding-top: 1rem !important;
        padding-left: 1.5rem;
        padding-right: 1.5rem;
    }
    [data-testid="stHeader"] { height: 2.2rem; background: transparent; }
    @media (max-width: 640px) {
        [data-testid="stMainBlockContainer"] {
            max-width: 100%;
            padding-left: 0.8rem;
            padding-right: 0.8rem;
            padding-top: 0.6rem !important;
        }
    }

    /* --- Пояснителни/вторични текстове: по-светъл цвят --- */
    [data-testid="stCaptionContainer"] p { color: #e3cccc !important; }

    /* --- Табове като "pill" бутони --- */
    [data-testid="stTabs"] [data-baseweb="tab-list"] { gap: 10px; border-bottom: none; }
    [data-testid="stTabs"] [data-baseweb="tab-highlight"] { display: none; }
    [data-testid="stTabs"] [data-baseweb="tab-border"] { display: none; }
    [data-testid="stTab"] {
        background: rgba(255,255,255,0.06);
        border: 1px solid rgba(255,255,255,0.10);
        border-radius: 14px;
        padding: 10px 8px;
        justify-content: center;
    }
    [data-testid="stTab"] p { font-weight: 600; font-size: 0.95rem; }
    [data-testid="stTab"]:nth-of-type(1)[aria-selected="true"] {
        background: linear-gradient(135deg, #3b5bfd, #2540c9);
        border-color: transparent;
    }
    [data-testid="stTab"]:nth-of-type(2)[aria-selected="true"] {
        background: linear-gradient(135deg, #9333ea, #6b21a8);
        border-color: transparent;
    }

    /* --- Картови контейнери: st.container(key="card_...") --- */
    div[class*="st-key-card_"] {
        background: rgba(40, 10, 10, 0.55);
        border: 1px solid rgba(220, 70, 70, 0.30) !important;
        border-radius: 16px;
        padding: 14px 16px 4px 16px;
    }
    .bn-label { font-weight: 600; font-size: 0.98rem; margin-bottom: 2px; }
    .bn-sublabel { color: #e3cccc; font-size: 0.82rem; font-weight: 400; }

    /* --- Слайдер зелен (допълва темата), номер-инпут да пасва на картата --- */
    [data-testid="stNumberInput"] input {
        background: rgba(0,0,0,0.35) !important;
        border: 1px solid #22c55e !important;
        color: #4ade80 !important;
        font-weight: 700 !important;
        font-size: 1.25rem !important;
        border-radius: 10px !important;
        text-align: center;
    }
    [data-testid="stNumberInputStepDown"], [data-testid="stNumberInputStepUp"] { display: none; }
    /* Visually appended € sign (widget value itself stays numeric) */
    [data-testid="stNumberInputContainer"] { position: relative; }
    [data-testid="stNumberInputContainer"]::after {
        content: "€";
        position: absolute;
        right: 14px;
        top: 50%;
        transform: translateY(-50%);
        color: #4ade80;
        font-weight: 700;
        font-size: 1.1rem;
        pointer-events: none;
    }
    [data-testid="stNumberInputField"] { padding-right: 30px !important; }

    /* --- Стат карти: st.container(key="stat_...") --- */
    div[class*="st-key-stat_"] {
        background: rgba(40, 10, 10, 0.55);
        border: 1px solid rgba(220, 70, 70, 0.30) !important;
        border-radius: 16px;
        text-align: center;
        padding: 16px 10px !important;
    }
    div[class*="st-key-stat_"] [data-testid="stMarkdownContainer"] { width: 100%; text-align: center; margin-bottom: 0 !important; }
    .bn-stat-label { color: #e7d3d3; font-size: 0.85rem; margin-bottom: 4px; text-align: center; }
    .bn-stat-value { color: #4ade80; font-size: 2.3rem; font-weight: 800; line-height: 1.1; text-align: center; }
    .bn-stat-unit { color: #e3cccc; font-size: 0.8rem; text-align: center; }

    /* --- Tip box: st.container(key="tip_box") --- */
    div[class*="st-key-tip_box"] {
        background: rgba(34, 197, 94, 0.08);
        border: 1px solid rgba(74, 222, 128, 0.35) !important;
        border-radius: 14px;
        display: flex !important;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        min-height: 68px;
        padding: 14px 16px !important;
    }
    div[class*="st-key-tip_box"] [data-testid="stMarkdownContainer"] { width: 100%; text-align: center; margin-bottom: 0 !important; }
    div[class*="st-key-tip_box"] p { color: #86efac !important; margin: 0; text-align: center; }

    /* --- CTA бутони (зелени, като в дизайна) --- */
    [data-testid="stButton"] button, [data-testid="stLinkButton"] a {
        background: linear-gradient(135deg, #22c55e, #16a34a) !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: 14px !important;
        font-weight: 700 !important;
        padding: 0.8rem 1rem !important;
    }

    /* --- Бутон за изтегляне на доклада: отделен, забележим червен цвят --- */
    [data-testid="stDownloadButton"] button {
        background: linear-gradient(135deg, #ef4444, #b91c1c) !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: 14px !important;
        font-weight: 700 !important;
        padding: 0.8rem 1rem !important;
        box-shadow: 0 0 0 1px rgba(255, 255, 255, 0.15) inset;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# API Keys Check от Secrets
api_key = st.secrets.get("OPENAI_API_KEY", "")

# 📥 УЛАВЯНЕ НА ДАННИТЕ ОТ URL
session_id = st.query_params.get("session_id")

# 🧪 ТЕСТОВ РЕЖИМ, ВИДИМ САМО ЗА СОБСТВЕНИКА (за донастройки на живо, без риск
# обикновени посетители да ползват тестова карта вместо реално плащане).
#
# - Преди плащане: тестовият линк се показва само ако в URL-а има таен параметър
#   ?owner_test=<OWNER_TEST_TOKEN>, известен само на теб. Обикновените посетители
#   никога не виждат и не могат да отворят тестовия линк.
# - След плащане: ?owner_test не преживява Stripe redirect-а (Stripe връща само
#   session_id), затова режимът се определя directly от префикса на самото
#   session_id - Stripe маркира тестовите сесии с "cs_test_", реалните с "cs_live_".
#   Това работи независимо дали тайният параметър е останал в адреса.
STRIPE_API_KEY_LIVE = st.secrets.get("STRIPE_API_KEY_LIVE", st.secrets.get("STRIPE_API_KEY", ""))
STRIPE_API_KEY_TEST = st.secrets.get("STRIPE_API_KEY_TEST", "")
OWNER_TEST_TOKEN = st.secrets.get("OWNER_TEST_TOKEN", "")
APP_URL = st.secrets.get("APP_URL", "https://biz-navigator.streamlit.app")

if session_id:
    use_test_mode = session_id.startswith("cs_test_")
else:
    use_test_mode = bool(OWNER_TEST_TOKEN) and st.query_params.get("owner_test") == OWNER_TEST_TOKEN

stripe.api_key = STRIPE_API_KEY_TEST if use_test_mode else STRIPE_API_KEY_LIVE

if use_test_mode:
    st.sidebar.warning("🧪 Тестов режим е активен (видим само за теб чрез тайния параметър).")

# Базови стойности по подразбиране в session_state
if "fixed_costs" not in st.session_state: st.session_state.fixed_costs = 1200
if "price" not in st.session_state: st.session_state.price = 50
if "cost" not in st.session_state: st.session_state.cost = 20
if "idea_text" not in st.session_state: st.session_state.idea_text = ""
if "verified_sessions" not in st.session_state: st.session_state.verified_sessions = {}
if "generated_reports" not in st.session_state: st.session_state.generated_reports = {}

# 💾 ПОСТОЯННО СЪХРАНЕНИЕ НА ИДЕЯТА + ФИНАНСИТЕ ЧРЕЗ STRIPE METADATA
#
# По-рано идеята пътуваше кодирана в client_reference_id (лимит ~200 знака,
# само alphanumeric/-/_) - оттам идваше нуждата от съкращаване. Stripe
# metadata на Checkout Session позволява до 50 полета по 500 знака всяко,
# с произволни символи (включително кирилица), затова разбиваме дългата идея
# на "чънкове" от по IDEA_CHUNK_SIZE знака в отделни metadata полета -
# практически премахва ограничението (до MAX_IDEA_CHUNKS * IDEA_CHUNK_SIZE
# знака общо). Данните се записват директно в Stripe при създаването на
# Checkout Session (виж по-долу) и се препрочитат надеждно след плащането -
# без нужда от отделна база данни.
IDEA_CHUNK_SIZE = 480
MAX_IDEA_CHUNKS = 20  # 20 * 480 = 9600 знака таван, далеч над всякаква реална идея

def build_report_metadata(fixed_costs, price, cost, idea):
    idea = idea[: IDEA_CHUNK_SIZE * MAX_IDEA_CHUNKS]
    chunks = [idea[i:i + IDEA_CHUNK_SIZE] for i in range(0, len(idea), IDEA_CHUNK_SIZE)] or [""]
    metadata = {
        "fc": str(fixed_costs),
        "price": str(price),
        "cost": str(cost),
        "idea_chunks": str(len(chunks)),
    }
    for i, chunk in enumerate(chunks):
        metadata[f"idea_{i}"] = chunk
    return metadata


def read_report_metadata(metadata):
    if metadata is None:
        return None
    try:
        # В по-новите версии на stripe-python "metadata" е StripeObject, който
        # вече НЕ поддържа .get(...) като обикновен dict (хвърля AttributeError) -
        # затова първо го конвертираме в чист Python dict.
        data = metadata.to_dict() if hasattr(metadata, "to_dict") else dict(metadata)
        if not data:
            return None
        n_chunks = int(data.get("idea_chunks", "0") or "0")
        idea = "".join(data.get(f"idea_{i}", "") for i in range(n_chunks))
        return {
            "fixed_costs": int(data.get("fc", 0)),
            "price": int(data.get("price", 0)),
            "cost": int(data.get("cost", 0)),
            "idea": idea,
        }
    except Exception as e:
        print(f"[Stripe metadata] Неуспешно прочитане: {e}")
        return None


def get_stripe_session(sid):
    if not sid:
        return None
    try:
        return stripe.checkout.Session.retrieve(sid)
    except Exception as e:
        print(f"[Stripe] Грешка при извличане на сесия {sid}: {e}")
        return None

stripe_session = None
is_payment_valid = False

if session_id:
    if session_id in st.session_state.verified_sessions:
        stripe_session = st.session_state.verified_sessions[session_id]
    else:
        with st.spinner("🔒 Проверка на статуса на плащането..."):
            stripe_session = get_stripe_session(session_id)
        st.session_state.verified_sessions[session_id] = stripe_session

    if stripe_session is not None:
        is_payment_valid = stripe_session.payment_status == "paid"
    else:
        st.sidebar.warning("⚠️ Неуспешна проверка на плащането. Презаредете страницата или се свържете с поддръжка.")

if is_payment_valid and st.session_state.get("applied_payment_session") != session_id:
    decoded = read_report_metadata(stripe_session.metadata if stripe_session else None)
    if decoded:
        st.session_state.fixed_costs = decoded["fixed_costs"]
        st.session_state.price = decoded["price"]
        st.session_state.cost = decoded["cost"]
        st.session_state.idea_text = decoded["idea"]
    st.session_state.applied_payment_session = session_id

def add_formatted_paragraph(doc, text):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def create_docx(business_idea, financials, report_text):
    doc = Document()
    doc.add_heading("БИЗНЕС НАВИГАТОР - ЕКСПЕРТЕН ДОКЛАД", level=0)
    doc.add_heading("Обща информация", level=1)
    doc.add_paragraph(f"Бизнес идея: {business_idea}")
    doc.add_paragraph(f"Финансови параметри:\n"
                      f"- Постоянни месечни разходи: {financials['fc']} евро\n"
                      f"- Продажна цена: {financials['pr']} евро\n"
                      f"- Себестойност: {financials['cs']} евро\n"
                      f"- Нужни продажби за нулата: {financials['be']} бр./месец")
    doc.add_paragraph("\n" + "="*40 + "\n")
    doc.add_heading("Подробен анализ и план за действие", level=1)

    for paragraph in report_text.split("\n"):
        if paragraph.strip():
            if paragraph.strip().startswith("###"):
                clean_title = paragraph.replace("###", "").strip()
                doc.add_heading(clean_title, level=2)
            else:
                add_formatted_paragraph(doc, paragraph.strip())

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def synced_metric(key, label_html, min_v, max_v, step, help_text=None):
    """Слайдер + число, синхронизирани в двете посоки, увити в стилизирана карта."""
    slider_key = f"{key}__slider"
    num_key = f"{key}__num"
    if slider_key not in st.session_state:
        st.session_state[slider_key] = st.session_state[key]
    if num_key not in st.session_state:
        st.session_state[num_key] = st.session_state[key]

    def _from_slider():
        v = st.session_state[slider_key]
        st.session_state[key] = v
        st.session_state[num_key] = v

    def _from_num():
        v = max(min_v, min(max_v, st.session_state[num_key]))
        st.session_state[key] = v
        st.session_state[slider_key] = v
        st.session_state[num_key] = v

    with st.container(border=True, key=f"card_{key}"):
        st.markdown(f'<div class="bn-label">{label_html}</div>', unsafe_allow_html=True)
        col_slider, col_num = st.columns([2.4, 1], vertical_alignment="center")
        with col_slider:
            st.slider("", min_value=min_v, max_value=max_v, step=step,
                       key=slider_key, on_change=_from_slider, label_visibility="collapsed",
                       help=help_text)
        with col_num:
            st.number_input("", min_value=min_v, max_value=max_v, step=step,
                             key=num_key, on_change=_from_num, label_visibility="collapsed")
    return st.session_state[key]


st.title("🚀 Бизнес Навигатор")
st.caption("Твоят дигитален стартъп ментор")

if is_payment_valid:
    st.balloons()
    st.success("🎉 Плащането е потвърдено успешно! Твоят персонализиран бизнес план се генерира...")

    current_idea = st.session_state.idea_text
    if not current_idea or len(current_idea) < 5:
        current_idea = "Бизнес модел на база финансови калкулации."

    if not api_key:
        st.error("🔑 Липсва OpenAI API ключ в настройките.")
    else:
        margin = st.session_state.price - st.session_state.cost
        be_units = math.ceil(st.session_state.fixed_costs / margin) if margin > 0 else 0
        min_turnover = be_units * st.session_state.price

        if session_id in st.session_state.generated_reports:
            full_report = st.session_state.generated_reports[session_id]
        else:
            full_report = None
            with st.spinner("🤖 Експертният AI консултант анализира ТВОЯТА идея..."):
                try:
                    client = OpenAI(api_key=api_key)

                    daily_sales_needed = be_units / 30 if be_units else 0

                    paid_prompt = f"""
                    Ти си топ бизнес консултант с 20 години опит в съветването на стартиращи
                    компании в България — работил си с банки, инвеститори и над 500 малки бизнеса.
                    Пишеш доклади, за които клиентите казват, че си им спестил месеци проучване.

                    Напиши изключително подробен, дълбоко персонализиран и максимално професионален
                    бизнес доклад на български език. Никакви общи приказки, никакви клишета като
                    "усърдна работа" или "вярвай в себе си" — само конкретни, приложими съвети,
                    основани на числата и идеята на потребителя. Пиши все едно клиентът е платил
                    500 евро за консултация с реален експерт, не за автоматичен текст. Не се
                    ограничавай в дължината - целта е максимална дълбочина и полезност, не краткост.
                    Всяка точка по-долу трябва да съдържа конкретни разсъждения и примери, не само
                    едно изречение - обясни ЗАЩО, не само КАКВО.

                    КОНКРЕТНА ИДЕЯ НА ПОТРЕБИТЕЛЯ: "{current_idea}"

                    ВЪВЕДЕНИ ФИНАНСИ ОТ ПОТРЕБИТЕЛЯ:
                    - Постоянни месечни разходи: {st.session_state.fixed_costs} евро
                    - Продажна цена на бройка/час: {st.session_state.price} евро
                    - Себестойност на бройка/минимален разход: {st.session_state.cost} евро
                    - Марж на бройка/час: {margin} евро
                    - Точка на баланса (Break-even): Нужни са точно {be_units} продажби на месец
                      (средно {daily_sales_needed:.1f} на ден) за оборот от {min_turnover} евро,
                      само за да покрият разходите (нулева печалба).

                    Докладът ТРЯБВА да съдържа следните раздели, с точно тези заглавия, в този ред,
                    и да е богат на конкретика — навсякъде, където е relevant, цитирай директно
                    числата по-горе:

                    ### 📊 ЧАСТ 1: ОБОБЩЕНИЕ ЗА РЪКОВОДСТВОТО (EXECUTIVE SUMMARY)
                    - 4-5 изречения, обобщаващи реалистичната перспектива пред този конкретен бизнес:
                      колко трудно/лесно е да се стигне до break-even с тези числа, какъв е
                      най-големият шанс за успех, и какъв е най-големият риск, ако не се действа
                      целенасочено.

                    ### 💪 ЧАСТ 2: ПЕРСОНАЛИЗИРАН АНАЛИЗ НА МОДЕЛА
                    - СИЛНИ СТРАНИ: 3-те най-големи стратегически предимства на този конкретен модел,
                      всяко обяснено в 3-4 изречения защо точно за тази идея и тези числа е предимство,
                      с конкретен пример как да се използва на практика.
                    - СКРИТИ РИСКОВЕ: 3-те най-големи опасности специално за този бизнес на българския
                      пазар (данъци и осигуровки, скрита/наситена конкуренция, регулации, сезонност,
                      кешфлоу), с КОНКРЕТЕН, изпълним начин за избягване на всеки риск - не просто
                      описание на проблема.

                    ### 🎯 ЧАСТ 3: ЦЕЛЕВА АУДИТОРИЯ И КОНКУРЕНТЕН ПЕЙЗАЖ
                    - Профил на идеалния клиент за тази идея (демография, поведение, къде се намира).
                      Защо точно той/тя ще плати за това.
                    - 2-3 вероятни конкурента (по тип, не измислени имена) на българския пазар и
                      конкретно с какво тази идея може да се разграничи и да победи на всеки от тях.

                    ### 💰 ЧАСТ 4: ФИНАНСОВА СТРАТЕГИЯ И ЦЕНООБРАЗУВАНЕ
                    - Подробна интерпретация какво точно означава маржът от {margin} евро на бройка/час
                      за устойчивостта на бизнеса и за колко бързо се натрупва буфер за неочаквани разходи.
                    - Конкретна препоръка: трябва ли цената/себестойността да се коригират, и с колко,
                      за да бъде break-even точката реалистично достижима по-бързо.
                    - Три сценария за първите 6 месеца (консервативен, реалистичен, оптимистичен) —
                      за всеки: очакван брой продажби на месец, прогнозна печалба/загуба в евро,
                      и в кой месец се очаква реален кумулативен break-even.

                    ### 🗺️ ЧАСТ 5: ПЪТНА КАРТА ПО СТЪПКИ (ROADMAP)
                    - Детайлна хронологична пътна карта, специфична за тази идея, разбита по:
                      Седмица 1, Седмица 2, Седмица 3-4, Месец 2, Месец 3, Месец 4-6 — какво точно
                      се прави във всяка от тези фази, защо е точно в тази последователност, и какъв
                      измерим резултат трябва да има на края на всяка фаза.

                    ### 📣 ЧАСТ 6: МАРКЕТИНГ И ПРИДОБИВАНЕ НА ПЪРВИ КЛИЕНТИ
                    - 3 конкретни, нискобюджетни маркетингови канала/тактики, подходящи специално
                      за тази идея и българския пазар, с приблизителен бюджет, очакван резултат и
                      конкретна първа стъпка за стартиране на всеки канал.

                    ### 📝 ЧАСТ 7: AI ЧЕК-ЛИСТ СЪС ЗАДАЧИ ЗА ТАЗИ СЕДМИЦА
                    - Списък от точно 5 конкретни, практически задачи, специфични за този бизнес,
                      подредени по приоритет, които потребителят може да изпълни веднага, всяка с
                      кратко обяснение защо е важна точно сега.

                    ### 🎯 ЗАКЛЮЧЕНИЕ
                    - 3-4 изречения с честна, но мотивираща обобщена оценка, и ЕДНО единствено
                      най-важно следващо действие, което потребителят трябва да направи утре.

                    Форматирай с ясни markdown заглавия и bullet points където е подходящо.
                    Бъди максимално конкретен, цитирай цифрите навсякъде, където е relevant,
                    и никъде не изпадай в общи приказки, приложими за произволен бизнес.
                    """

                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{"role": "user", "content": paid_prompt}],
                        temperature=0.7,
                        max_tokens=4096,
                    )

                    full_report = response.choices[0].message.content
                    st.session_state.generated_reports[session_id] = full_report

                except Exception as e:
                    print(f"[OpenAI] Грешка при генериране на доклад за сесия {session_id}: {e}")
                    st.error("⚠️ Възникна временен проблем при генерирането на доклада. Моля, презаредете страницата.")

        if full_report:
            st.markdown("## 📊 ТВОЯТ ПЕРСОНАЛИЗИРАН ЕКСПЕРТЕН БИЗНЕС ПЛАН")
            st.markdown(full_report)
            st.markdown("---")

            fin_data = {
                "fc": st.session_state.fixed_costs,
                "pr": st.session_state.price,
                "cs": st.session_state.cost,
                "be": be_units
            }

            docx_bytes = create_docx(current_idea, fin_data, full_report)

            st.download_button(
                label="📥 Изтегли Бизнес плана в редактируем Word (.docx) формат",
                data=docx_bytes,
                file_name="Business_Plan_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

st.markdown("---")

# Tabs
tab1, tab2 = st.tabs(["📊 1. Сметни риск", "💡 2. AI Валидация"])

# TAB 1: Financial Simulator
with tab1:
    st.subheader("📱 Финансов симулатор")
    st.caption("Нагласи слайдерите или въведи стойност на ръка, за да видиш минимума за оцеляване:")

    synced_metric(
        "fixed_costs",
        "💼 Месечни постоянни разходи<br><span class='bn-sublabel'>(наем, осигуровки)</span>",
        min_v=200, max_v=10000, step=100,
    )
    synced_metric(
        "price",
        "💰 Продажна цена за 1 бройка / час",
        min_v=0, max_v=500, step=1,
    )
    synced_metric(
        "cost",
        "📦 Себестойност на 1 бройка<br><span class='bn-sublabel'>(материали/доставка)</span>",
        min_v=0, max_v=300, step=1,
    )

    if st.session_state.price <= st.session_state.cost:
        st.error("🛑 Цената трябва да е по-висока от себестойността!")
    else:
        margin = st.session_state.price - st.session_state.cost
        be_units = math.ceil(st.session_state.fixed_costs / margin)
        min_turnover = be_units * st.session_state.price
        daily_sales = be_units / 30

        st.markdown("#### Резултат за твоя бизнес:")
        col1, col2 = st.columns(2)
        with col1:
            with st.container(border=True, key="stat_sales"):
                st.markdown(
                    f'<div class="bn-stat-label">Нужни продажби</div>'
                    f'<div class="bn-stat-value">{be_units}</div>'
                    f'<div class="bn-stat-unit">бр./мес.</div>',
                    unsafe_allow_html=True,
                )
        with col2:
            with st.container(border=True, key="stat_turnover"):
                st.markdown(
                    f'<div class="bn-stat-label">Минимум оборот</div>'
                    f'<div class="bn-stat-value">{min_turnover}</div>'
                    f'<div class="bn-stat-unit">евро</div>',
                    unsafe_allow_html=True,
                )

        with st.container(border=True, key="tip_box"):
            st.markdown(
                f"💡 Това означава средно по **{daily_sales:.1f}** продажби на ден, за да излезеш на нула."
            )

# TAB 2: AI Validation
with tab2:
    st.subheader("🤖 Запиши идеята си")
    text_idea = st.text_area(
        "Напиши или коригирай идеята си тук:",
        placeholder="Пример: Искам да отворя автомивка с 4 бокса в центъра на София...",
        key="idea_text",
        max_chars=3000,
        help="Можеш да пишеш свободно и по-подробно - идеята се пази изцяло (без ограничение на практика) и се използва цялата за персонализирания доклад.",
    )

    if st.button("🚀 Анализирай моята идея", use_container_width=True):
        if not api_key:
            st.error("🔑 Липсва OpenAI API ключ.")
        elif len(text_idea) < 5:
            st.warning("⚠️ Моля, въведете описание.")
        else:
            with st.spinner("AI Менторът изчислява финансовия риск..."):
                try:
                    client = OpenAI(api_key=api_key)
                    free_prompt = f"Ти си бизнес консултант. Направи КРАТЪК преглед (до 3 изречения) на идея: '{text_idea}'."
                    response = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": free_prompt}], temperature=0.7)

                    st.markdown("---")
                    st.markdown("### 🔓 Твоят безплатен предварителен анализ:")
                    st.info(response.choices[0].message.content)

                    st.markdown("### 📊 Отключи Пълния Експертен Доклад")

                    checkout_url = None
                    try:
                        metadata = build_report_metadata(
                            st.session_state.fixed_costs,
                            st.session_state.price,
                            st.session_state.cost,
                            text_idea,
                        )
                        checkout_session = stripe.checkout.Session.create(
                            mode="payment",
                            line_items=[{
                                "price_data": {
                                    "currency": "eur",
                                    "product_data": {"name": "Бизнес Навигатор — Пълен Експертен Доклад"},
                                    "unit_amount": 499,
                                },
                                "quantity": 1,
                            }],
                            metadata=metadata,
                            success_url=f"{APP_URL}/?session_id={{CHECKOUT_SESSION_ID}}",
                            cancel_url=APP_URL,
                        )
                        checkout_url = checkout_session.url
                    except Exception as e:
                        print(f"[Stripe] Грешка при създаване на checkout сесия: {e}")
                        st.error("⚠️ Възникна проблем при подготовката на плащането. Моля, опитайте отново.")

                    if checkout_url:
                        if use_test_mode:
                            st.caption("🧪 Тестов режим — картата 4242 4242 4242 4242 не таксува нищо реално.")

                        st.write("Нашият AI ще състави подробна пътна карта и чек-лист специално за тези стойности.")
                        st.link_button("💳 Отключи Пълния Бизнес Доклад за 4.99 евро", checkout_url, use_container_width=True)
                except Exception as e:
                    print(f"[OpenAI] Грешка при безплатен анализ: {e}")
                    st.error("⚠️ Възникна временен проблем. Моля, опитайте отново.")
